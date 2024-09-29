import pdfplumber as pb
import docx
from docx import Document

doc = docx.Document()
paragraph3 = doc.add_paragraph()
# 读取PDF文档
pdf = pb.open("C:/Users/Chen Junkai/Desktop/20074315.pdf")

# 获取页数
a = len(pdf.pages)
print("当前页：", a)
print("-----------------------------------------")

i = 0
for i in range(0, a):
    first_page = pdf.pages[i]
    print("本页：", first_page.page_number)
    print("-----------------------------------------")

    # 导出当前页文本
    text = first_page.extract_text()
    paragraph3.add_run(text)

doc.save("C:/Users/Chen Junkai/Desktop/输出结果.docx")
