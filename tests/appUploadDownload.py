import os
import time
import flaskr
import logging
import re
import difflib
import docx
#from win32com.client import Dispatch
from pycorrector.macbert.macbert_corrector import MacBertCorrector
from flask import Flask, render_template, request, send_from_directory

app = Flask(__name__)
def set_log(app):
    print('set_log start')
    gunicorn_logger = logging.getLogger('gunicorn.error')
    print(f'gunicorn_logger.handlers: {gunicorn_logger.handlers}')
    print(f'gunicorn_logger.level: {gunicorn_logger.level}')

    app.logger.handlers = gunicorn_logger.handlers
    app.logger.setLevel(gunicorn_logger.level)
    print(f'app.logger.handlers: {app.logger.handlers}')
    print(f'app.logger.level: {app.logger.level}')
    print(f'app.name: {app.name}')

    print('set_log end')
set_log(app)

# before_request 登录校验白名单
white_list = ['/login.html', '/demo.html', '/static', '/upload.html']

m = MacBertCorrector()

@app.before_request
def before_request():
    # app.logger.info(f"before_request. {request.url}")
    #print('print 1 before_request')
    #app.logger.debug('debug before_request')
    #app.logger.info('debug before_request')
    #app.logger.error('error before_request')

    # if request.path in white_list:
    #     return None
    for white in white_list:
        if request.path.startswith(white):
            return None

@app.route('/')
def hello_world():
    return 'Hello World!'


@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    # 渲染文件
    return render_template('upload.html')


# 文件保存的目录，根据实际情况的文件结构做调整；
# 若不指定目录，可以写成f.save(f.filename)，可以默认保存到当前文件夹下的根目录
# 设置上传文件保存路径 可以是指定绝对路径，也可以是相对路径（测试过）
app.config['UPLOAD_FOLDER'] = './upload'	## 该目录需要自行创建
# 将地址赋值给变量
file_dir = app.config['UPLOAD_FOLDER']


@app.route('/uploader', methods=['GET', 'POST'])
def uploader():
    """  文件上传  """
    if request.method == 'POST':
        # input标签中的name的属性值
        f = request.files['file']

        # 拼接地址，上传地址，f.filename：直接获取文件名
        f.save(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))
        # 输出上传的文件名
        print(request.files, f.filename)
        do_check(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))

        file_handle = open(os.path.join(app.config['UPLOAD_FOLDER'], f.filename), 'r')

        '''
        @after_this_request
        def remove_file(response):
            try:
                os.remove(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))
                file_handle.close()
            except Exception as error:
                app.logger.error("Error removing or closing downloaded file handle", error)
            return response

        return send_file(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))
        '''

        return send_from_directory(app.config['UPLOAD_FOLDER'], f.filename, as_attachment=True)
        return 'ok!'
    else:
        return render_template('upload.html')


def do_check(filename):

    check_result = 'check_result'
    #d = 'd'
    '''上方是无用部分,下面是读取doc文本'''
    document = docx.Document(filename)
    all_paragraphs = document.paragraphs
    #file_addr = filename
    correctsent, originwrongsent = correct(filename)
    print('!!!!!!!!!!!!')
    print(correctsent)
    print(originwrongsent)
    print('!!!!!!!!!!!!!')
    # 获取所有段落
    # print(type(all_paragraphs))
    for paragraph in all_paragraphs:
        for errstr in originwrongsent:
            print(type(errstr))
            if errstr in paragraph.text:
                i = originwrongsent.index(errstr)
                comment = paragraph.add_comment((''.join(correctsent[i])), 'rob')
    document.save(filename)

    '''下方是win批注部分
    correctsent, originwrongsent = correct(inputStringVer)#v2.0,都是list类型，originwrongsent用于遍历找到错误句子位置，并用correctsent批注
    word = Dispatch('Word.Application')
    word.Visible = 0
    # path = 'testdocinput.docx'
    # os.path.abspath('C:/Users/Administrator/Desktop/1.jpg')
    doc = word.Documents.Open(os.path.abspath('/data/pycorrector/testdocinput.docx'), Encoding='gbk')
    # doc = word.Documents.Open(FileName=path, Encoding='gbk')
    str_list = originwrongsent
    i = 0
    for err in str_list:
        correctstring =''.join(correctsent[i])
        text_find = err
        s = word.Selection
        s.Start = 0
        s.End = 0
        s.Find.Execute(text_find)
        doc.Comments.Add(Range=word.Selection.Range, Text=correctstring)
        i = i+1

    '''
    #check_result = "\n".join(check_result_list)    #v1.0,无用信息

    # check pwd
    if True:
        if check_result == "":
            check_result = "==> 未检测到错误 ==<"
        #flash("static/test210.html")
        #flash(html_stringver)
        '''
        with open("flaskr/static/test210.html", "r", encoding='utf8') as f:
            html_str = f.read()
            print(type(html_str))
            flash(html_str)
        '''
        # abort(400)
        #return render_template('test210.html') org
        #return render_template('demo.html', last_input=input_text)
    #print(f"2. check_result: {check_result}")
    # app.logger.info(f"修改密码成功. username: {user}")
    return 0

def correct(file_addr):
    global m
    document = docx.Document(file_addr)
    all_paragraphs = document.paragraphs    #v2.0
    correct_part = []
    err_part = []
    #orgtxt = [] #v1.0 difflib所需
    #cortxt = [] #同上
    check_result_list = [] #v0.0
    for paragraph in all_paragraphs:
        # 打印每一个段落的文字
        print(paragraph.text)
        '''test'''
        # error_sentences = input_text.split("\n")
        error_sentences = re.split(r"[\n，]", paragraph.text)
        for line in error_sentences:
            print("error_sentences:{}\n".format(line))
            if line is None or line == "":
                print('empty line. skip')
                continue
            correct_sent, err = m.macbert_correct(line)
            if len(err) == 0:
                continue
            print("\n\nquery:{}\n   => {}\n  err:{}".format(line, correct_sent, err))
            correct_part.append(correct_sent)
            err_part.append(line)
            #check_result_list.append("=> {}\n  err:{}".format(correct_sent, err))

    '''
    if not input_text:
        return check_result_list
    '''
    '''
    #error_sentences = input_text.split("\n")
    error_sentences = re.split(r"[\n。！？]",input_text)

    for line in error_sentences:
        print("error_sentences:{}\n".format(line))
        if line is None or line == "":
            print('empty line. skip')
            continue
        correct_sent, err = m.macbert_correct(line)
        if len(err) == 0:
            continue
        print("\n\nquery:{}\n   => {}\n  err:{}".format(line, correct_sent, err))
        check_result_list.append("=> {}\n  err:{}".format(correct_sent, err))
    '''

    #下方是v1.0,为了用difflib显示可视化结果，因此完全显示两段文本并写入到html字符串中
    '''
    error_sentences = re.split(r"[\n，]", input_text)
    res = m.batch_macbert_correct(error_sentences)
    for sent, r in zip(error_sentences, res):
        #print("original sentence final:{} => {} err:{}".format(sent, r[0], r[1])) org
        check_result_list.append("original sentence final:{} => {} err:{}".format(sent, r[0], r[1]))
        orgtxt.append("，{}".format(sent))
        cortxt.append("，{}".format(r[0]))
    #with open("flaskr/static/test210.html", "w", encoding='utf8') as f:
    txt1 = "".join(orgtxt)
    txt2 = "".join(cortxt)
    txt1_lines = txt1.splitlines()
    txt2_lines = txt2.splitlines()
    d = difflib.HtmlDiff(wrapcolumn=20)
    #f.write(d.make_file(txt1_lines, txt2_lines))
    d = d.make_file(txt1_lines, txt2_lines)
    #webbrowser.open("flaskr/test210.html")
    '''
    return correct_part, err_part


"""
    运行项目
"""
if __name__ == '__main__':
   app.run(debug=True, host='0.0.0.0', port=5001)
