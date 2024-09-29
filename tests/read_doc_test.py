import docx
import os
from win32com.client import Dispatch
#from docx import Document
errlist = ['总经理', '沦海当前，必当扬帆破浪，任重道远，更需砥砺前行']
correctlist = ['总经历', '沧海当前']
# Open the Word document
document = docx.Document("C:/Users/Chen Junkai/Desktop/testdocinput.docx")
all_paragraphs = document.paragraphs
# Find the text to add a comment to
for paragraph in all_paragraphs:
    for errstr in errlist:
        if errstr in paragraph.text:
            print(type(paragraph.text))
            print(type(errstr))
            i = errlist.index(errstr)
            comment = paragraph.add_comment((''.join(correctlist[i])), '在此，我代表中粮资本班子')
            '''
    if '总经理' or '沦海当前，必当扬帆破浪，任重道远，更需砥砺前行' in paragraph.text:
        comment = paragraph.add_comment('Your comment here','在此，我代表中粮资本班子')
        run = paragraph.add_run('text1')
        '''
# Save the document
document.save("C:/Users/Chen Junkai/Desktop/testdocinput.docx")


'''
document = docx.Document("C:/Users/Chen Junkai/Desktop/testdocinput.docx")
inputstringver = ''
#获取所有段落
all_paragraphs = document.paragraphs
#print(type(all_paragraphs))
for paragraph in all_paragraphs:
    #打印每一个段落的文字
    #print(paragraph.text)
    inputstringver = inputstringver + paragraph.text
    #print(type(paragraph.text))
print(inputstringver)
'''

'''
word = Dispatch('Word.Application')
word.Visible = 0
#path = 'testdocinput.docx'
#os.path.abspath('C:/Users/Administrator/Desktop/1.jpg')
doc = word.Documents.Open(os.path.abspath('C:/Users/Chen Junkai/Desktop/testdocinput.docx'), Encoding='gbk')
#doc = word.Documents.Open(FileName=path, Encoding='gbk')
str_list = ['在此，我代表中粮资本班子', '这一年，国际国内经济形势错综复杂']
#print(str_list[0])
i = 0
for err in str_list:
    stringver =''.join(str_list[i])
    text_find = err
    s = word.Selection
    s.Start = 0
    s.End = 0
    s.Find.Execute(text_find)#s.Find.Execute(text_find[0])
    doc.Comments.Add(Range=word.Selection.Range, Text=stringver)
    i = i+1
'''
'''
指定单句注释
find_t = r'2021年，是中粮资本“十四五”发展的开局之年'
word.Selection.Find.Execute(find_t)
doc.Comments.Add(Range=word.Selection.Range, Text='测试1')
'''
#指定par批注
#doc.Comments.Add(Range=doc.paragraphs[1].Range, Text='测试')