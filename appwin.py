import random

import chardet
from aip import AipOcr

from flask import request, abort, redirect, render_template, session, make_response, g, flash, send_from_directory
from markupsafe import escape
import flaskr
import logging
import re
import pythoncom
import difflib
import docx
import os
import win32com
import pdfplumber as pb
from shutil import copyfile
from win32com.client import Dispatch
from pycorrector.macbert.macbert_corrector import MacBertCorrector
from UserSecretFile import UserSecretFile
import pymysql
from flask_mail import Mail, Message
import time
import sys
from threading import Timer, Thread, BoundedSemaphore
from ruamel.yaml import YAML
from docx.shared import Pt
import jionlp as jio
import json


captcha_global = ''
email_global = ''
time_start = 0
captcha_en = True
count = 60
timer_en = False

global captcha_stat_global # 表示当前captcha的生成状态, 1为captcha有效, 0为captcha无效
captcha_stat_global = 0

filenum = 0
page = 0
now_stage = 0 #当前所处的状态（初始化为default）
pre_stage = 0 #上一个状态（初始化为default）

# 信号量相关
Max = 1
s = BoundedSemaphore(Max) # 用信号量控制'demo2.html'的AJAX与uploader()的执行顺序
s.acquire()

s_demo2 = BoundedSemaphore(Max) # 用信号量控制'demo1.html'的AJAX与do_Check()的执行顺序
s_demo2.acquire()

app = flaskr.create_app()
mail = Mail(app)

def _send_async_mail(app, message):
    with app.app_context():
        mail.send(message)

def send_mail(subject, to, body):
    '''异步发送邮件的代码，目前该函数无法与视图函数联合处理不存在邮件地址的问题'''
    message = Message(subject, recipients=to, body=body)
    thr = Thread(target=_send_async_mail, args=[app, message])
    thr.start()
    return thr

def countdown():
    global count, timer_en, captcha_en
    if timer_en:
        if count > 0:
            count = count - 1
        else:
            count = 60
            captcha_en = True
            timer_en = False
    Timer(1, countdown).start()
tr = Timer(1, countdown)
tr.start()

def set_log(app):
    print('set_log start')
    gunicorn_logger = logging.getLogger('gunicorn.error')
    print(f'gunicorn_logger.handlers: {gunicorn_logger.handlers}')
    print(f'gunicorn_logger.level: {gunicorn_logger.level}')

    app.logger.handlers = gunicorn_logger.handlers
    app.logger.setLevel(gunicorn_logger.level)
    print(f'app.logger.handlers: {app.logger.handlers}')
    print(f'app.logger.level: {app.logger.level}')
    print(f'app.name: {app.name}')

    print('set_log end')


db = pymysql.connect(host='localhost', user='root', password='leju123', database='correction')
cursor = db.cursor()
set_log(app)
userFile = UserSecretFile(app.config["PWD_FILE"])
white_list = ['/login.html', '/login', '/static', '/forgetpassword.html', '/forgetpassword', '/resetpassword.html', '/resetpassword', '/setemail.html', '/setemail', '/changepassword.html', '/do_changepwd', '/changeemail.html', '/do_changeemail']
home_page = '/home'
home_page = '/demo.html'

@app.route('/submit', methods=['POST'])
def submit():
  select_value = request.form.get('models')
  # 处理数据
  global now_stage,pre_stage
  if (select_value == 'default'):
      g.flag = 0 #这个不能删，这是作为demo中判断哪个要显示的
      pre_stage = now_stage #继承上一个状态
      now_stage = 0 #新的状态是default
  elif(select_value == 'train'):
      g.flag = 1
      pre_stage = now_stage  # 继承上一个状态
      now_stage = 1  # 新的状态是default
  print(pre_stage)
  print(now_stage)
  # flash("切换模型成功") # 已通过js的alert()函数实现同样目的

  #界面跳转
  if(page == 0):
      return render_template('demo.html')
  elif (page == 1):
      return render_template('demo1.html')
  elif (page == 2):
      return render_template('demo2.html')
  elif (page == 3):
      return render_template('demo3.html')
  else:
      return render_template('demo.html')

@app.route('/turn1', methods=['POST'])
def turn1():
    global page
    page = 1
    return render_template('demo1.html')

@app.route('/turn2', methods=['POST'])
def turn2():
    global page
    page = 2
    return render_template('demo2.html')

@app.route('/turn3', methods=['POST'])
def turn3():
    global page
    page = 3
    return render_template('demo3.html')

@app.route('/turn_back', methods=['POST'])
def turn_back():
    global page
    global now_stage, pre_stage, filenum
    global filenum
    page = 0
    now_stage = 0  # 重新都变成默认的
    pre_stage = 0
    if ((filenum % 2) == 1):  # 奇数，说明是不好的那个模型
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                filenum + 1) + '.bin')
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                filenum) + '.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin')
        filenum = filenum + 1
    return render_template('demo.html')



@app.before_request
def before_request():
    for white in white_list:
        if request.path.startswith(white):
            return None

    if 'auth_ok' not in session or 'username' not in session: # 只有"auth_ok"和"username"均设置后才能进入'login.html'
        app.logger.info(f"before_request. need login.")
        res = make_response(redirect('login.html'))
        return res
    else:
        # 已登录，设置全局变量
        g.user = {'username': session.get('username')}

@app.route("/")
@app.route("/demo.html", methods=['GET'])
def index():
    return render_template('demo.html')

@app.route("/login.html", methods=['GET'])
def login_page():
    return render_template('login.html')

@app.route("/forgetpassword.html", methods=['GET'])
def forgetpassword_page():
    return render_template('forgetpassword.html')

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    # 渲染文件
    return render_template('upload.html')

file_dir = app.config['UPLOAD_FOLDER']


def remove_substr(str,sub):
   if str.endswith(sub):
      return str[:-len(sub)]
   return str

global filedir_stat
filedir_stat = 0 # 标识'demo2.html'中是否上传了文件, 0为未上传图片, 1为上传图片

@app.route('/function_Get_filedir_stat', methods=['GET'])
def function_Get_filedir_stat():
    # print("AJAX calls func...")
    print("filedir_stat: "+str(filedir_stat))
    s.acquire()
    print("exit AJAX func...")
    return json.dumps(filedir_stat)

@app.route('/uploader', methods=['GET', 'POST'])
def uploader():
    """  文件上传  """
    if request.method == 'POST':
        # input标签中的name的属性值
        f = request.files['file']
        tmpfileName = str(time.time()).split(".")[0]#根据时间戳生成一个新文件名
        tmpfileName = tmpfileName+'.docx'
        global filedir_stat, page
        if f.filename == '': # 上传文件为空
            filedir_stat = 1
            s.release()
            if page == 2:
                return render_template('demo2.html')
            elif page == 3:
                return render_template('demo3.html')
            else:
                return render_template('demo.html')
            # return render_template('demo2_emptyFileDir.html')
        else: # 上传文件不为空
            # 拼接地址，上传地址，f.filename：直接获取文件名
            f.save(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))

        # 输出上传的文件名
        #print(request.files, f.filename)
        # 上传文件格式错误
        if (f.filename.endswith("docx") == False and f.filename.endswith("pdf") == False and f.filename.endswith("jpg") == False and f.filename.endswith("txt") == False):
            print('文件类型错误')
            filedir_stat = -1
            s.release()
            os.remove(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))#不符合类型的文件不予保留
            # 判断是不是只有一个管理员用户
            if page == 2:
                return render_template('demo2.html')
            elif page == 3:
                return render_template('demo3.html')
            else:
                return render_template('demo.html')
        else: # 上传文件格式正确, 即上传文件格式一定是.txt, .doc, .docx, .jpg, .pdf之中的一种
            if page == 3 and f.filename.endswith("txt") is True:
                # filedir_stat = 0
                # s.release()
                print("准备生成数据集")
                trainmodel(os.path.join(app.config['UPLOAD_FOLDER'], f.filename),
                           os.path.join("D:/pycorrector/pycorrector/macbert/output/", remove_substr(f.filename, ".txt"),
                                        ".json"))
                if filedir_stat == -2:
                    return render_template('demo3.html')
                else:
                    return render_template('demo3_train_finish.html')
            elif page == 2:
                if f.filename.endswith("txt") is True:
                    filedir_stat = -1
                    s.release()
                    return render_template('demo2.html')
                else: # 除去.txt格式, 剩下的文件格式都是合法的
                    filedir_stat = 0
                    s.release()
                    if(f.filename.endswith("pdf") == True):
                        pdf2docx(os.path.join(app.config['UPLOAD_FOLDER'], f.filename), os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))
                        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))  # 删除用户原先上传的文件
                        do_check_doc(os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))  # 对时间戳命名的新文件进行检查并批注
                        flash("pdf纠错中")
                        return send_from_directory(app.config['UPLOAD_FOLDER'], tmpfileName, as_attachment=True)  # 给用户返回并下载新生成的批注文件
                    elif(f.filename.endswith("jpg") == True):
                        jpg2word(os.path.join(app.config['UPLOAD_FOLDER'], f.filename), os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))
                        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))  # 删除用户原先上传的文件
                        do_check_doc(os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))  # 对时间戳命名的新文件进行检查并批注
                        flash("图片纠错中")
                        return send_from_directory(app.config['UPLOAD_FOLDER'], tmpfileName, as_attachment=True)  # 给用户返回并下载新生成的批注文件
                    else:
                        copyfile(os.path.join(app.config['UPLOAD_FOLDER'], f.filename), os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))#将用户上传的原文件复制到根据时间戳命名的文件中
                        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))#删除用户原先上传的文件
                        do_check_doc(os.path.join(app.config['UPLOAD_FOLDER'], tmpfileName))#对时间戳命名的新文件进行检查并批注
                        flash("word纠错中")
                        return send_from_directory(app.config['UPLOAD_FOLDER'], tmpfileName, as_attachment=True)#给用户返回并下载新生成的批注文件
            else:
                filedir_stat = -1
                s.release()
                if page == 2:
                    return render_template('demo2.html')
                elif page == 3:
                    return render_template('demo3.html')
                else:
                    return render_template('demo.html')
    #else:
        #return render_template('upload.html')

def pdf2docx(filename,tmpfilename):
    doc = docx.Document()
    paragraph3 = doc.add_paragraph()
    # 读取PDF文档
    pdf = pb.open(filename)
    # 获取页数
    a = len(pdf.pages)
    print('页数'+str(a))
    i = 0
    for i in range(0, a):
        first_page = pdf.pages[i]
        # 导出当前页文本
        text = first_page.extract_text()
        paragraph3.add_run(text)
    doc.save(tmpfilename)
    pdf.close()
    return 0


def do_check_doc(filename):
    check_result = 'check_result'
    correctsent, originwrongsent = correct_doc(filename)#v2.0,都是list类型，originwrongsent用于遍历找到错误句子位置，并用correctsent批注
    pythoncom.CoInitialize()    #dispatch必须要使用多线程模块
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = 0    #后台运行
    word.DisplayAlerts = 0  # 不显示，不警告
    doc = word.Documents.Open(filename, Encoding='gbk')
    str_list = originwrongsent
    i = 0
    for err in str_list:
        correctstring =''.join(correctsent[i])
        print('改正'+correctstring)
        text_find = err
        s = word.Selection
        s.Start = 0
        s.End = 0
        s.Find.Execute(text_find)
        print('对应位置'+text_find)
        doc.Comments.Add(Range=word.Selection.Range, Text=correctstring)
        i = i+1
    doc.Save()
    doc.Close()
    word.Quit()  # 关闭 office
    pythoncom.CoUninitialize()#关闭多线程

    # check pwd
    if True:
        if check_result == "":
            check_result = "==> 未检测到错误 ==<"
        return render_template('demo.html')
    print(f"2. check_result: {check_result}")
    # app.logger.info(f"修改密码成功. username: {user}")
    return res


def correct_doc(filename):
    global filenum
    if pre_stage == now_stage:
        m1 = MacBertCorrector('pytorch_model.bin')
    else:
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                filenum + 1) + '.bin')
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                filenum) + '.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin')
        filenum = filenum + 1
        m1 = MacBertCorrector('pytorch_model.bin')
    document = docx.Document(filename)
    all_paragraphs = document.paragraphs
    correct_part = []#修改后的正确句子
    err_part = []#原文错误句子
    for paragraph in all_paragraphs:
        # 打印每一个段落的文字
        print(paragraph.text)
        error_sentences = re.split(r"[\n，]", paragraph.text)#将原文进行分割
        for line in error_sentences:
            print("error_sentences:{}\n".format(line))
            if line is None or line == "":
                print('empty line. skip')
                continue
            correct_sent, err = m1.macbert_correct(line)#correct_sent是修正之后的句子
            if len(err) == 0:
                continue
            print("\n\nquery:{}\n   => {}\n  err:{}".format(line, correct_sent, err))
            correct_part.append(correct_sent)#将修正之后的句子传到correct_part 这个list中
            err_part.append(line)#将相应的原文中的错误句子传到err_part中

    return correct_part, err_part

global textempty_stat
textempty_stat = 0 # 标识'demo1.html'中文本框是否为空, 0为空, 1为有内容

@app.route('/function_Get_textarea_stat', methods=['GET'])
def function_Get_textarea_stat():
    # print("AJAX calls func textarea...")
    print("textempty_stat: " + str(textempty_stat))
    s_demo2.acquire()
    print("exit AJAX func textarea...")
    return json.dumps(textempty_stat)

'''do_check函数是用于文本纠错，从文本框中传入需要查错的文本'''
@app.route("/do_check", methods=['POST'])
def do_check():
    input_text = request.form['input_text']
    print(f"1. input_text: {input_text}")
    global textempty_stat
    if input_text == '' or len(input_text) > 1000:
        textempty_stat = 1
        s_demo2.release()
        global now_stage, pre_stage, filenum
        global filenum
        now_stage = 0  # 重新都变成默认的
        pre_stage = 0
        if ((filenum % 2) == 1):  # 奇数，说明是不好的那个模型
            os.rename(
                'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin',
                'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                    filenum + 1) + '.bin')
            os.rename(
                'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                    filenum) + '.bin',
                'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin')
            filenum = filenum + 1
        return render_template('demo1.html')

    textempty_stat = 0
    s_demo2.release()
    check_result_list, html_stringver = correct(input_text)#第一个返回值没有用处，html_stringver是经过对比生成的difflib的html结果的string类型文本

    check_result = "\n".join(check_result_list)

    # check pwd
    if True:
        if check_result == "":
            check_result = "==> 未检测到错误 ==<"
        flash(html_stringver)#将difflib的html结果传递到demo.html中用来显示
        if (now_stage == 0):
            g.flag = 0
        elif (now_stage == 1):
            g.flag = 1
        return render_template('demo1.html', last_input=input_text)
    print(f"2. check_result: {check_result}")
    # app.logger.info(f"修改密码成功. username: {user}")
    return res

def correct(input_text):
    global filenum
    if pre_stage == now_stage:
        m = MacBertCorrector('pytorch_model.bin')
    else:
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(
                filenum+1) + '.bin')
        os.rename(
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model' + str(filenum) + '.bin',
            'C:/Users/Tinky\'s_Comoputer/.pycorrector/datasets/macbert_models/chinese_finetuned_correction/pytorch_model.bin')
        filenum = filenum + 1
        m = MacBertCorrector('pytorch_model.bin')
    orgtxt = []
    cortxt = []
    check_result_list = []
    if not input_text:
        return check_result_list
    error_sentences = re.split(r"[\n，]", input_text)
    res = m.batch_macbert_correct(error_sentences)
    for sent, r in zip(error_sentences, res):
        check_result_list.append("original sentence final:{} => {} err:{}".format(sent, r[0], r[1]))
        orgtxt.append("，{}".format(sent))#代表了原文的句子内容
        cortxt.append("，{}".format(r[0]))#修改后的句子内容
    txt1 = "".join(orgtxt)
    txt1 = txt1[1:]#去除开头的‘，’
    txt2 = "".join(cortxt)
    txt2 = txt2[1:]#去除开头的‘，’
    txt1_lines = txt1.splitlines()
    txt2_lines = txt2.splitlines()
    d = difflib.HtmlDiff(wrapcolumn=20)#创建html文件，wrapcolumn是代表每行的字符数多少，用于控制显示的html文件
    d = d.make_file(txt1_lines, txt2_lines)#根据两段话生成对应的difflib html结果，并用字符串的形式传回
    return check_result_list, d




def check_user_pwd(username,password):
    sql = """SELECT * FROM users WHERE usersname=\'""" + username + '\''
    try:
        cursor.execute(sql)
        results = cursor.fetchall()

        if results == (): # 用户不存在
            app.logger.error(f"login. username: {escape(username)} not found")
            return None
    except:
        return None
    password_inDB = results[0][1]
    email = results[0][2]
    numOfLogIn = results[0][3]
    global email_global
    email_global = email
    if password_inDB != password: # 密码不对
        app.logger.error(f"login. username: {escape(username)} invalid password")
        return None
    if numOfLogIn == 0:
        return 2
    return 1

def check_password_rule(password):
    # err_msg = '密码规则： 4~20字符，不能包含冒号、回车换行'
    # length 4 ~ 20
    if len(password) < 4 or len(password) > 20:
        return False
    # not contain \n :
    if ':' in password or '\n' in password:
        return False
    return True

@app.route("/login", methods=['POST'])
def login():
    username = request.form['username']
    password = request.form['password']
    app.logger.info(f"login. 1 username: {escape(username)} password: {escape(password)}")

    otp_user = check_user_pwd(username, password)

    if otp_user is None:
        flash('用户名或密码错误', category="error")
        # abort(403)
        return render_template('login.html')
    if otp_user == 2: # 用户此前没设置过邮箱, 是第一次登陆
        # session['auth_ok'] = True
        session['username'] = username
        print(f"登录成功. username: {username}")
        app.logger.info(f"登录成功. username: {username}")
        flash('首次登陆请设置邮箱', category="info")
        return render_template('setemail.html')
    session['auth_ok'] = True
    session['username'] = username
    print(f"登录成功. username: {username}")
    app.logger.info(f"登录成功. username: {username}")
    res = make_response(redirect(home_page))
    return res

def check_emailformat(email):
    '''通过正则表达式检查邮箱格式是否有问题'''
    regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    if re.fullmatch(regex, email):
        return True
    else:
        return False

@app.route("/setemail",methods=['POST'])
def setemail():
    user = session['username']
    newemail = request.form.get('newemail') # 目前用session['temp_newemail']保证第一次设置的邮箱一定是收到验证码的邮箱
    bt_assure = request.values.get("确认")
    bt_send = request.values.get("发送验证码")
    global captcha_global, captcha_stat_global, captcha_en, count, timer_en
    if bt_assure == '确认':
        captcha_input = request.form['captcha']
        if captcha_input == captcha_global and captcha_stat_global:
            sql_setnewemail = """UPDATE users SET email=\'""" + session.get('temp_newemail') + '\' WHERE usersname=\'' + user + '\''
            sql_setnumOfLogIn = """UPDATE users SET numOfLogIn=""" + str(1) + ' WHERE usersname=\'' + user + '\''
            session.pop('temp_newemail')
            captcha_stat_global = 0
            try:
                # 执行sql语句
                cursor.execute(sql_setnumOfLogIn) # 修改"numOfLogIn"
                cursor.execute(sql_setnewemail)   # 设置"email"
                # 提交到数据库执行
                db.commit()
                flash("电子邮箱设置成功", category="info")
                global email_global
                email_global = newemail
                session['auth_ok'] = True
                res = make_response(redirect(home_page))
                return res
            except:
                # 如果发生错误则回滚
                db.rollback() # "numOfLogIn"和"email"必须均修改成功才算完成对邮箱的设置
                flash('设置失败', category="error")
                return render_template('setemail.html')
        else:
            flash('验证码错误', category="error")
            return render_template('setemail.html')
    if bt_send == '发送验证码':
        if not captcha_en:
            flash('请等待' + str(count) + '秒', category="warning")
        elif check_emailformat(newemail): # 能到这里, captcha_en一定有效
            captcha_en = False
            timer_en = True
            rcp = []
            rcp.append(newemail)
            msg = Message('您的验证码', recipients=rcp)
            captcha_global = generate_random()
            msg.body = '欢迎使用快准-公文纠错系统\n您的验证码是：' + captcha_global
            try:
                mail.send(msg)
            except Exception as e:
                print(e)
                flash('电子邮箱地址无效', category="error")
                return render_template('setemail.html')
            captcha_stat_global = 1
            session['temp_newemail'] = newemail
            flash('验证码已发送', category="info")
        else:
            flash('请输入有效的电子邮箱地址', category="error")
        return render_template('setemail.html')

@app.route("/forgetpasswordhtml",methods=['POST'])
def forgetpasswordhtml():
    return render_template('forgetpassword.html')


def validate_email(email):
    '''检查邮箱是否存在于数据库中, 并获取相应的用户名到session['username']中'''
    sql = """SELECT * FROM users WHERE email =\'""" + email + '\''
    try:
        cursor.execute(sql)
        results = cursor.fetchall()

        if results == ():
            app.logger.error(f"forgetpassword. email: {escape(email)} not found")
            return False
    except:
        return False
    session['username'] = results[0][0]
    return True

@app.route("/forgetpassword",methods=['POST'])
def forgetpassword():
    bt_assure = request.values.get("确认")
    bt_send = request.values.get("发送验证码")
    bt_back = request.values.get("返回")

    global captcha_global, captcha_stat_global, captcha_en, count, timer_en

    if bt_assure == '确认':
        captcha_input = request.form['captcha']
        if captcha_input == captcha_global and captcha_stat_global:
            captcha_stat_global = 0  # 只要点击确认后, 本次生成的captcha就设置为无效, 无论是否通过captcha检验, 下次再点击确认前都需要再次发送邮件
            return render_template('resetpassword.html')
        else:
            flash('验证码无效', category="error")
            return render_template('forgetpassword.html')
    if bt_send == '发送验证码':
        username = request.form['username']
        sql = """SELECT * FROM users WHERE usersname=\'""" + username + '\''
        try:
            cursor.execute(sql)
            results = cursor.fetchall()

            if results == ():  # 用户不存在
                flash('用户不存在', category='error')
                return render_template('forgetpassword.html')
        except:
            return None
        cemail = results[0][2]
        recipient = request.form['email_address']
        if cemail != recipient:
            flash('与预留邮箱不相同', category='error')
            return render_template('forgetpassword.html')
        if validate_email(recipient) is False:
            flash('请输入您设置的电子邮箱地址', category="error")
        elif captcha_en:
            captcha_stat_global = 1
            captcha_en = False
            timer_en = True
            rcp = []
            rcp.append(recipient)
            msg = Message('您的验证码', recipients=rcp)
            captcha_global = generate_random()
            msg.body = '欢迎使用快准-公文纠错系统\n您的验证码是：' + captcha_global
            try:
                mail.send(msg) # 此处似乎没必要进行try except, 因为此处只会向数据库中存在的邮箱发邮件, 数据库中的邮箱几乎可以肯定是有效的
            except Exception as e:
                print(e)
                flash('电子邮箱地址无效', category="error")
                return render_template('forgetpassword.html')
            flash('验证码已发送', category="info")
        else: # 邮箱存在于数据库中但captcha_en无效
            flash('请等待' + str(count) + '秒', category="warning")
        return render_template('forgetpassword.html')
    if bt_back == '返回':
        captcha_stat_global = 0
        if session.get('username') is not None:
            session.pop('username')
        return render_template('login.html')




def generate_random():
    return ''.join(random.sample('0123456789', 6))



@app.route("/resetpassword",methods=['POST'])
def resetpassword():
    # username = request.form['username']
    username = session['username'] # 用户名信息来自'/forgetpassword'的session
    password = request.form['newpassword']
    password_confirm = request.form['assurepassword']
    if password != password_confirm:
        flash('两次密码不一致', category="error")
        return render_template('resetpassword.html')
    if not check_password_rule(password):
        flash('密码规则： 4~20字符，不能包含冒号、回车换行', category="error")
        return render_template('resetpassword.html')
    sql = """UPDATE users SET password=\'""" + password + '\' WHERE usersname=\'' + username + '\''
    try:
        # 执行sql语句
        cursor.execute(sql)
        # 提交到数据库执行
        db.commit()
        flash("重置密码成功", category="info")
        return render_template('login.html')
    except:
        # 如果发生错误则回滚
        db.rollback()
        flash('error', category="error")
        return render_template('resetpassword.html')





@app.route("/logout", methods=['GET'])
def logout():
    username = session.get('username')
    del session['auth_ok']
    del session['username']
    res = make_response(redirect(home_page))
    res.set_cookie('newemail', '', expires=0)
    res.set_cookie('email_address', '', expires=0)
    res.set_cookie('username', '', expires=0)
    app.logger.info(f"退出登录. username: {username}")
    return res

@app.route("/changeemail", methods=['GET'])
def goto_changeemail():
    return render_template('changeemail.html')

@app.route("/changepassword", methods=['GET'])
def goto_changepwd():
    return render_template('changepassword.html')


@app.route("/do_changepwd", methods=['POST'])
def do_changepwd():
    if request.values.get("返回") == "返回":
        res = make_response(redirect(home_page))
        return res
    user = session['username']
    old_password = request.form['currentpassword']
    new_password = request.form['newpassword']
    new2_password = request.form['assurepassword']

    # app.logger.info(f"changepwd. 1 user: {user} old_password: {escape(old_password)} new_password: {escape(new_password)} new2_password: {escape(new2_password)}!")

    # check pwd
    if new_password != new2_password:
        flash('原密码错误或新密码不一致', category="error")
        # abort(400)
        return render_template('changepassword.html')

    otp_user = check_user_pwd(user, old_password)
    if otp_user is None:
        flash('原密码错误或新密码不一致', category="error")
        # abort(403)
        return render_template('changepassword.html')

    if not check_password_rule(new_password):
        flash('密码规则： 4~20字符，不能包含冒号、回车换行', category="error")
        return render_template('changepassword.html')

    # app.logger.info(f"exec_change_pwd. 1 otpFile.show")
    # otpFile.show()

    sql = """UPDATE users SET password=\'""" + new_password + '\' WHERE usersname=\'' + user + '\''
    try:
        # 执行sql语句
        cursor.execute(sql)
        # 提交到数据库执行
        db.commit()
        flash('修改密码成功', category="info")
        res = make_response(redirect(home_page))
        print(f"修改密码成功. username: {user}")
        app.logger.info(f"修改密码成功. username: {user}")
        return res
    except:
        # 如果发生错误则回滚
        db.rollback()
        flash('修改失败', category="error")
        return render_template('changepassword.html')
    # otpFile.show()


@app.route("/do_changeemail", methods=['POST'])
def do_changeemail():
    user = session['username']
    newemail = request.form.get('newemail') # 如果'newemail'为空, 返回的应该是None
    bt_assure = request.values.get("确认")
    bt_send = request.values.get("发送验证码")
    bt_back = request.values.get("返回")
    global captcha_global, captcha_stat_global, captcha_en, count, timer_en
    if bt_assure == '确认':
        captcha_input = request.form['captcha']
        if captcha_input == captcha_global and captcha_stat_global:
            captcha_stat_global = 0
            # sql = """UPDATE users SET email=\'""" + newemail + '\' WHERE usersname=\'' + user + '\''
            sql = """UPDATE users SET email=\'""" + session.get('temp_newemail') + '\' WHERE usersname=\'' + user + '\''
            session.pop('temp_newemail')
            try:
                # 执行sql语句
                cursor.execute(sql)
                # 提交到数据库执行
                db.commit()
                flash('修改邮箱成功', category="info")
                res = make_response(redirect(home_page))
                return res
            except:
                # 如果发生错误则回滚
                db.rollback()
                flash('修改邮箱失败', category="error")
                return render_template('changeemail.html')
        else:
            flash('验证码错误', category="error")
            return render_template('changeemail.html')
    if bt_send == '发送验证码':
        if not captcha_en:
            flash('请等待' + str(count) + '秒', category="warning")
        elif check_emailformat(newemail):
            captcha_en = False
            timer_en = True
            rcp = []
            rcp.append(newemail)
            msg = Message('您的验证码', recipients=rcp)
            captcha_global = generate_random()
            msg.body = '欢迎使用快准-公文纠错系统\n您的验证码是：' + captcha_global
            try:
                mail.send(msg)
                # send_mail('您的验证码', rcp, captcha_global)
            except Exception as e:
                print(e)
                flash('电子邮箱地址无效', category="error")
                return render_template('changeemail.html')
            captcha_stat_global = 1
            session['temp_newemail'] = newemail # 目前用session['temp_newemail']保证更新后的邮箱一定是收到验证码的邮箱
            flash('验证码已发送', category="info")
        else:
            flash('请输入有效的电子邮箱地址', category="error")
        return render_template('changeemail.html')
    if bt_back == '返回':
        captcha_stat_global = 0
        res = make_response(redirect(home_page))
        return res

def get_file_content(filePath):
    with open(filePath,'rb') as fp:
        return fp.read()

def jpg2word(filename,tmpfilename):
    APP_ID = '35592034'
    APP_KEY = 'ldizEtXSXGhHyGs0D8HWZ1Q9'
    SECRET_KEY = '915gyNyd5jI6MtVG3mhz5ve12BIBy3rm'
    client = AipOcr(APP_ID, APP_KEY, SECRET_KEY)
    image = get_file_content(filename)  # 获取图片
    msg = client.general(image)  # 调用API解析图片后生成一段信息储存在变量中
    print(msg)
    doc = docx.Document()
    for i in msg.get('words_result'):
        print(i.get('words'))
        doc.add_paragraph(i.get('words'))
    doc.save(tmpfilename)

def creatDataset_sentence(stro,rst):
    '''
    rest = jio.remove_parentheses(stro)
    rest = ''.join(rest)
    res = jio.homophone_substitution(rest)
    '''
    res = jio.homophone_substitution(stro)  #调用jionlp包生成混淆的同音词，默认生成三个list类型的错误样本
    res1 = res[:1]  #res1-3是提取三个混淆同音词的错误样本
    res2 = res[1:2]
    res3 = res[2:3]
    stro = str(stro)    #将数据转换成str格式
    str1 = ''.join(res1)
    str2 = ''.join(res2)
    str3 = ''.join(res3)

    temp1 = []
    if str1 == "":
        pass
    else:
        for i in range(len(str1)-1):        #找到与原文不匹配的部分，认为是错误部分，将其添加到rst中
            if stro[i] != str1[i]:
                temp1.append(i)

        rst.append({'id': "-",
                    'original_text': str1,
                    'wrong_ids': temp1,
                    'correct_text': stro})
    temp2 = []
    if str2 == "":
        pass
    else:
        for i in range(len(str2)-1):
            if stro[i] != str2[i]:
                temp2.append(i)
        rst.append({'id': "-",
                    'original_text': str2,
                    'wrong_ids': temp2,
                    'correct_text': stro})

    temp3 = []
    if str3 == "":
        pass
    else:
        for i in range(len(str3)-1):
            if stro[i] != str3[i]:
                temp3.append(i)
        rst.append({'id': "-",
                    'original_text': str3,
                    'wrong_ids': temp3,
                    'correct_text': stro})

def trainmodel(filename,resultfile):
    rst = []
    print(type(rst))
    global filedir_stat
    with open(filename, 'rb') as f:
        if 'GB2312' != chardet.detect(f.read())['encoding']:
            filedir_stat = -2
            s.release()
            return # 不执行后续代码

    filedir_stat = 0
    s.release()
    with open(filename) as f:  # 原始的txt数据
        read_data = f.read()
        read_data = read_data.split("\n")  # 去掉回车的\n
        read_data = "".join(read_data)
        a = re.split(r"[。！]", read_data)  # 用。以及！切割长字符串
        for strof in a:
            if strof == '' or strof[0] == "（" or "0" <= strof[0] <= "9" or len(strof) <= 10 or len(
                    strof) >= 80:  # 筛选传入的字符，如果开头带有（，数字认为样本不够干净，去除这句话
                continue  # 同时长度大于80也会去除该句话，因为过长会导致训练出错
            creatDataset_sentence(strof, rst)  # 生成错误样本的list，结果在rst中
    rst = json.dumps(rst, indent=4, ensure_ascii=False)

    with open('D:/pycorrector/pycorrector/macbert/output/train74.json', 'r+', encoding='utf8') as f:  # 输入需要添加数据集的.json文件
        strtemp = f.read()

        if strtemp == '':
            rst = json.dumps(rst, indent=4, ensure_ascii=False)

        else:
            content = json.loads(strtemp)
            content = json.dumps(content, indent=4, ensure_ascii=False)  # 将list转为str

            f.seek(0)  # 指向文本开头，这两行的用处是因为r+只能读写原文的内容并向后追加，加入后就能实现更改前两行的内容
            f.truncate()  # 清除文本

            '''此部分是因为rst和content必须转变为str才能写入，就会存在原文件最后一行多余的]以及新文件第一行多余的[，并且中间少了，连接'''
            content = content[:len(content) - 2]
            content = content + ","
            rst = rst[1:]
            rst = content + rst
        f.write(rst)
        print("success!")

    yaml = YAML()
    yaml.preserve_quotes = True
    with open('D:/pycorrector/pycorrector/macbert/train_macbert4csc.yml') as f:
        test = yaml.load(f)
    test['DATASETS']['TRAIN'] = "output/train74.json"
    with open("D:/pycorrector/pycorrector/macbert/train_macbert4csc.yml", 'w') as file:
        yaml.dump(test, file)

    os.system('python D:/pycorrector/pycorrector/macbert/train.py')

if __name__ == '__main__':
   app.run(debug=True, host='0.0.0.0', port=5001)
